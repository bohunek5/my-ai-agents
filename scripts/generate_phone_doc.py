from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import qn
from docx.oxml import OxmlElement
import os

def set_cell_background(cell, fill_color):
    shading_elm = cell._element.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), fill_color)
    shading_elm.append(shading)

def create_phone_table(container_cell, data):
    # Tworzymy tabelę wewnątrz komórki
    table = container_cell.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    # Maksymalna szerokość dla połówki A4 (ok. 10cm)
    table.columns[0].width = Cm(7.5)
    table.columns[1].width = Cm(2.5)
    
    # Nagłówek tabeli
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Osoba / Dział'
    hdr_cells[1].text = 'Nr'
    
    for cell in hdr_cells:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0] if p.runs else p.add_run()
            run.bold = True
            run.font.size = Pt(12)

    for section, employees in data:
        row = table.add_row().cells
        row[0].text = section
        row[1].text = ""
        set_cell_background(row[0], 'D9D9D9') # Ciemniejszy szary dla kontrastu
        for p in row[0].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(11)

        for name, ext in employees:
            row = table.add_row().cells
            row[0].text = name
            row[1].text = ext
            # Wszystko pogrubione i duże
            for p in row[0].paragraphs:
                for r in p.runs: 
                    r.bold = True
                    r.font.size = Pt(11)
            for p in row[1].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs: 
                    r.bold = True
                    r.font.size = Pt(11)

def generate_4x_phone_list():
    doc = Document()
    
    # Ultra małe marginesy (0.5 cm)
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Cm(0.5)
        section.right_margin = Cm(0.5)

    data = [
        ("ZARZĄD", [("Krzysztof Bara", "61"), ("Radosław Narwojsz", "62")]),
        ("SEKRETARIAT", [("Sylwia Suska", "25")]),
        ("KSIĘGOWOŚĆ", [("Agnieszka Bara", "52"), ("Karolina Laskowska", "51")]),
        ("MAGAZYNY", [("Magazyn Wydań", "71"), ("Montaż", "81"), ("Produkcja", "80")]),
        ("PRESCOT SP. Z O.O.", [("Natalia Lis", "31"), ("Adam Garbowski", "34"), ("Dariusz Nita", "35"), ("Anna Galor", "36"), ("Anna Asztemborska", "37")]),
        ("DZIAŁ GRAFICZNY", [("Kinga Bohdanowicz", "91"), ("Karol Bohdanowicz", "92")]),
        ("PRESCOT S.C.", [("Jarosław Badowski", "21"), ("Iwona Baczewska", "22"), ("Marcin Górniewicz", "24")])
    ]
    
    # Główna tabela układu 2x2 rozciągnięta na całą stronę
    layout_table = doc.add_table(rows=2, cols=2)
    # Szerokość A4 to 21cm, minus marginesy 1cm = 20cm (2 kolumny po 10cm)
    for col in layout_table.columns:
        col.width = Cm(10.0)
    
    for row_idx in range(2):
        for col_idx in range(2):
            cell = layout_table.cell(row_idx, col_idx)
            # Wstawiamy tylko tabelę (bez nagłówka tekstowego powyżej)
            create_phone_table(cell, data)

    output_path = "/Users/karolbohdanowicz/Downloads/Numery_Wewnetrzne_Prescot_A4x4.docx"
    doc.save(output_path)
    return output_path

    output_path = "/Users/karolbohdanowicz/Downloads/Numery_Wewnetrzne_Prescot_A4x4.docx"
    doc.save(output_path)
    return output_path

if __name__ == "__main__":
    path = generate_4x_phone_list()
    print(f"PLIK GENEROWANY (4x na A4): {path}")
