from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_training_doc():
    doc = Document()
    
    # Styl tytułu
    title = doc.add_heading('PROGRAM SZKOLENIOWY DLA NOWEGO HANDLOWCA - PRESCOT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("Poniższy dokument zawiera 3 kluczowe plany wdrożeniowe (scenariusze), które pozwalają na kompleksowe przeszkolenie nowej osoby z zakresu produktów, systemów i strategii sprzedaży w ekosystemie Prescot.")

    # --- PLAN 1 ---
    doc.add_heading('PLAN 1: "TECHNICZNY EKSPERT" (Scenariusz Głębokiej Wiedzy)', level=1)
    doc.add_paragraph("Cel: Zrozumienie specyfikacji technicznej i niezawodności komponentów.")
    
    table1 = doc.add_table(rows=1, cols=2)
    table1.style = 'Table Grid'
    hdr_cells = table1.rows[0].cells
    hdr_cells[0].text = 'Moduł'
    hdr_cells[1].text = 'Kluczowe zagadnienia'
    
    items1 = [
        ("TAŚMY LED (Fundament)", "Delux (budżet), Premium PL (produkcja w Polsce, powtarzalność barw), Premium COB (bez efektu kropek, 3 lata gwarancji)."),
        ("ZASILANIE (Serce Systemu)", "Scharfer (High-End, 5-7 lat gwarancji), Prescot (optymalizacja kosztów)."),
        ("PROFILE ALUMINIOWE", "Systemy Kluś – dobór osłon, zaślepek, sposoby montażu (radiatory)."),
        ("STEROWANIE", "Mi-Light / MiBoxer: Parowanie, strefy, piloty i aplikacje."),
        ("AKCESORIA I ŹRÓDŁA", "Złączki i przewodny Prescot (0.75mm), żarówki AR, świetlówki LED.")
    ]
    
    for mod, desc in items1:
        row = table1.add_row().cells
        row[0].text = mod
        row[1].text = desc

    doc.add_page_break()

    # --- PLAN 2 ---
    doc.add_heading('PLAN 2: "PROJEKTOWY ARCHITEKT" (Scenariusz Systemowy)', level=1)
    doc.add_paragraph("Cel: Nauka składania gotowych zestawów z dostępnych podzespołów.")
    
    p2_steps = [
        "KROK 1: Wybór źródła (Taśmy) - Dopasowanie do budżetu i estetyki (Delux / Premium PL / COB).",
        "KROK 2: Dobór Oprawy (Profile) - Integracja z profilami Kluś (aspekt chłodzenia i designu).",
        "KROK 3: System Zasilania - Obliczanie mocy: (Metraż * Moc taśmy) * 1.2 (zapas 20%). Wybór 24V dla stabilności.",
        "KROK 4: Inteligencja (Mi-Light) - Dobór sterowników radiowych i strefowych.",
        "KROK 5: Wykończenie - Sprzedaż dodatkowa (akcesoria montażowe, oprawy punktowe AR)."
    ]
    for step in p2_steps:
        doc.add_paragraph(step, style='List Bullet')

    doc.add_heading('Technika TCO (Total Cost of Ownership)', level=2)
    doc.add_paragraph("Argument dla klienta: Tanie LED (wymiana co rok + koszty ekipy) vs Prescot (praca 7 lat). Inwestujesz raz, oszczędzasz na serwisie.")

    # --- PLAN 3 ---
    doc.add_heading('PLAN 3: "STRATEGICZNY HANDLOWIEC" (Psychologia i Narzędzia)', level=1)
    doc.add_paragraph("Cel: Maksymalizacja marży i wykorzystanie narzędzi wsparcia sprzedaży.")
    
    p3_points = [
        "Hierarchia Gwarancji: COB (3 lata), Scharfer (do 7 lat) - to nasze tarcze przed konkurencją.",
        "USP: Realne parametry (nasze 10W to 10W, a nie 6W na papierze).",
        "Narzędzia WWW: tasmaled.com.pl/tasmyled (język korzyści), prescot.pl (inspiracje).",
        "Standard B2B: Terminowość i pewność dostawy są ważniejsze niż najniższa cena."
    ]
    for point in p3_points:
        doc.add_paragraph(point, style='List Number')

    # --- NUMERY WEWNĘTRZNE ---
    doc.add_page_break()
    doc.add_heading('WEWNĘTRZNA LISTA NUMERÓW - PRESCOT', level=1)
    
    contacts_data = [
        ("ZARZĄD", [("Zarząd (ogólny)", "31"), ("Krzysztof Bara", "61"), ("Radosław Narwojsz", "62")]),
        ("SEKRETARIAT", [("Sylwia Suska", "25")]),
        ("KSIĘGOWOŚĆ", [("Agnieszka Bara", "52"), ("Karolina Laskowska", "51")]),
        ("MAGAZYNY", [("Magazyn Wydań", "71"), ("Montaż", "81"), ("Produkcja", "80")]),
        ("PRESCOT SP. Z O.O.", [("Adam Garbowski", "34"), ("Dariusz Nita", "35"), ("Anna Galor", "36"), ("Anna Asztemborska", "37")]),
        ("DZIAŁ GRAFICZNY", [("Kinga Bohdanowicz", "91"), ("Karol Bohdanowicz", "92")]),
        ("PRESCOT S.C.", [("Jarosław Badowski", "21"), ("Iwona Baczewska", "22"), ("Marcin Górniewicz", "24")])
    ]
    
    ext_table = doc.add_table(rows=1, cols=2)
    ext_table.style = 'Table Grid'
    hdr = ext_table.rows[0].cells
    hdr[0].text = 'Osoba / Dział'
    hdr[1].text = 'Nr Wewnętrzny'
    
    for dept, persons in contacts_data:
        # Nagłówek działu w kolumnie
        row = ext_table.add_row().cells
        row[0].text = dept
        row[0].paragraphs[0].runs[0].bold = True # Bold for department
        row[1].text = ""
        
        for name, ext in persons:
            row = ext_table.add_row().cells
            row[0].text = name
            row[1].text = ext

    doc.add_paragraph("\nStopka: Wygenerowano przez system Antigravity Prescot AI.")
    
    output_path = "output/Plan_Wdrozeniowy_Handlowca.docx"
    os.makedirs("output", exist_ok=True)
    doc.save(output_path)
    return output_path

if __name__ == "__main__":
    path = create_training_doc()
    print(f"DOKUMENT GENEROWANY: {path}")
